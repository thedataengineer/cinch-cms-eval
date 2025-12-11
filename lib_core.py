"""
CINCH CMS Evaluation Framework – Core Libraries
Reusable Python classes for evaluation, scoring, and reporting
"""

import json
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from pydantic import BaseModel


# ============================================================================
# ONTOLOGY MODULE
# ============================================================================

@dataclass
class Capability:
    """Represents a CMS capability dimension."""
    key: str
    label: str
    facets: List[str]
    scale: str  # e.g., "0-3"
    importance: str  # "critical", "high", "medium"


@dataclass
class UseCase:
    """Represents a business use case."""
    key: str
    label: str
    required_capabilities: Dict[str, int]  # capability_key -> min_level


@dataclass
class BusinessOutcome:
    """Represents a business outcome/KPI."""
    key: str
    label: str
    weight: float


class CMSOntology:
    """Encapsulates the CMS evaluation ontology."""
    
    def __init__(self, ontology_dict: Dict[str, Any]):
        self.raw = ontology_dict
        self.capabilities = {
            k: Capability(
                key=k,
                label=v["label"],
                facets=v["facets"],
                scale=v["scale"],
                importance=v["importance"]
            )
            for k, v in ontology_dict.get("capabilities", {}).items()
        }
        self.use_cases = {
            k: UseCase(
                key=k,
                label=v["label"],
                required_capabilities=v["required_capabilities"]
            )
            for k, v in ontology_dict.get("use_cases", {}).items()
        }
        self.business_outcomes = {
            k: BusinessOutcome(
                key=k,
                label=v["label"],
                weight=v["weight"]
            )
            for k, v in ontology_dict.get("business_outcomes", {}).items()
        }
    
    @classmethod
    def from_file(cls, filepath: str) -> "CMSOntology":
        """Load ontology from JSON file."""
        with open(filepath, "r") as f:
            data = json.load(f)
        return cls(data)
    
    def get_capability(self, key: str) -> Optional[Capability]:
        return self.capabilities.get(key)
    
    def get_use_case(self, key: str) -> Optional[UseCase]:
        return self.use_cases.get(key)
    
    def capability_keys(self) -> List[str]:
        return list(self.capabilities.keys())
    
    def use_case_keys(self) -> List[str]:
        return list(self.use_cases.keys())


# ============================================================================
# EVALUATION MODULE
# ============================================================================

class PlatformAssessment(BaseModel):
    """Structured assessment of a platform against ontology."""
    platform: str
    capability_scores: Dict[str, int]  # capability_key -> score (0-3)
    strengths: List[str]
    weaknesses: List[str]
    best_for_use_case: str
    overall_fit_score: float
    ai_generated: bool = False
    notes: Optional[str] = None
    
    def to_dict(self) -> Dict:
        return self.model_dump()


class PlatformEvaluator:
    """Evaluates platforms using LLM with structured outputs."""
    
    def __init__(self, provider=None):
        """
        Initialize the evaluator with an LLM provider.
        
        Args:
            provider: LLMProvider instance (OllamaProvider or AnthropicProvider).
                      If None, defaults to AnthropicProvider.
        """
        if provider is None:
            from llm_providers import get_provider
            provider = get_provider("anthropic")
        self.provider = provider
    
    def evaluate(
        self,
        platform_name: str,
        ontology: CMSOntology,
        context: Optional[str] = None
    ) -> PlatformAssessment:
        """
        Evaluate a platform against the ontology using the configured LLM provider.
        
        Args:
            platform_name: Name of the CMS platform
            ontology: CMSOntology instance
            context: Optional contextual information (e.g., vendor docs snippet)
        
        Returns:
            PlatformAssessment with capability scores and analysis
        """
        
        # Build prompt
        capability_names = ", ".join([
            f"{cap.label} (scale 0-3)"
            for cap in ontology.capabilities.values()
        ])
        
        prompt = f"""
You are a CMS evaluation expert. Assess the '{platform_name}' CMS platform.

Evaluate across these capabilities (scale 0-3, where 3 is excellent):
{capability_names}

CONTEXT:
{context or "Use your knowledge of " + platform_name + " from public documentation."}

CINCH REQUIREMENTS:
- 20K+ paid views/day, 6K-7K unique visitors
- Primary goal: improve conversion rates and drive enrollments
- Currently spread across 5 CMS (HubSpot, Liferay, Ion, Starmark, Surefire)
- Want to consolidate but accept 3-platform reality
- Avoid Sitecore-scale monolith, avoid lightweight tools
- Interested in headless/composable approach

Provide your assessment as JSON with:
1. capability_scores: Dict of capability_key -> score (0-3)
2. strengths: 3 key strengths for CINCH
3. weaknesses: 3 key weaknesses for CINCH
4. best_for_use_case: Which use case this platform fits best
5. overall_fit_score: 0.0-1.0 overall fit score
"""
        
        # Simplified schema for Ollama compatibility
        assessment_schema = {
            "type": "object",
            "properties": {
                "platform": {"type": "string"},
                "capability_scores": {
                    "type": "object"
                },
                "strengths": {
                    "type": "array",
                    "items": {"type": "string"}
                },
                "weaknesses": {
                    "type": "array",
                    "items": {"type": "string"}
                },
                "best_for_use_case": {"type": "string"},
                "overall_fit_score": {"type": "number"}
            },
            "required": [
                "platform",
                "capability_scores",
                "strengths",
                "weaknesses",
                "best_for_use_case",
                "overall_fit_score"
            ]
        }
        
        # Call LLM with structured output
        response = self.provider.chat(prompt, assessment_schema)
        response_json = response.content
        
        return PlatformAssessment(
            platform=platform_name,
            capability_scores=response_json.get("capability_scores", {}),
            strengths=response_json.get("strengths", []),
            weaknesses=response_json.get("weaknesses", []),
            best_for_use_case=response_json.get("best_for_use_case", ""),
            overall_fit_score=response_json.get("overall_fit_score", 0.0),
            ai_generated=True,
            notes=f"Generated via {response.provider}"
        )


# ============================================================================
# SCORING MODULE
# ============================================================================

class CapabilityScorer:
    """Scores platforms against use cases and business outcomes."""
    
    def __init__(self, ontology: CMSOntology):
        self.ontology = ontology
    
    def score_for_use_case(
        self,
        platform_assessment: PlatformAssessment,
        use_case_key: str
    ) -> float:
        """
        Score platform fit for a specific use case.
        
        Returns: 0.0-1.0 score
        """
        use_case = self.ontology.get_use_case(use_case_key)
        if not use_case:
            return 0.0
        
        # Check required capabilities
        total_fit = 0.0
        count = 0
        
        for cap_key, required_level in use_case.required_capabilities.items():
            actual_score = platform_assessment.capability_scores.get(cap_key, 0)
            # Fit is how close actual is to required
            fit = min(actual_score / max(required_level, 1), 1.0)
            total_fit += fit
            count += 1
        
        return total_fit / count if count > 0 else 0.0
    
    def composite_score(
        self,
        platform_assessment: PlatformAssessment,
        use_case_keys: List[str],
        outcome_weights: Dict[str, float] = None
    ) -> float:
        """
        Calculate composite score weighted by use cases and business outcomes.
        
        Args:
            platform_assessment: Assessment to score
            use_case_keys: List of use case keys to consider
            outcome_weights: Optional custom outcome weights
        
        Returns: 0.0-1.0 composite score
        """
        
        # Use case fit (average across selected use cases)
        use_case_scores = [
            self.score_for_use_case(platform_assessment, uc_key)
            for uc_key in use_case_keys
        ]
        use_case_fit = sum(use_case_scores) / len(use_case_scores) if use_case_scores else 0.0
        
        # Business outcome fit (based on overall platform score)
        business_fit = platform_assessment.overall_fit_score
        
        # Weighted composite
        return (use_case_fit * 0.6) + (business_fit * 0.4)


# ============================================================================
# REPORTER MODULE
# ============================================================================

class ReportGenerator:
    """Generates evaluation reports in various formats."""
    
    def generate_markdown(
        self,
        evaluations: List[PlatformAssessment],
        recommendations: List[str],
        title: str = "CMS Evaluation Report"
    ) -> str:
        """Generate markdown report."""
        
        lines = [
            f"# {title}",
            f"**Generated:** {datetime.now().isoformat()}",
            "",
            "## Executive Summary",
            "Evaluation of CMS platforms against business and technical requirements.",
            "",
            "## Platform Assessments",
            ""
        ]
        
        for eval in evaluations:
            lines.extend([
                f"### {eval.platform}",
                f"**Overall Fit:** {eval.overall_fit_score:.2f}/1.0",
                "",
                "**Strengths:**",
                *[f"- {s}" for s in eval.strengths],
                "",
                "**Weaknesses:**",
                *[f"- {w}" for w in eval.weaknesses],
                "",
                f"**Best for:** {eval.best_for_use_case}",
                ""
            ])
        
        lines.extend([
            "## Recommendations",
            "",
            *[f"**{r}**" for r in recommendations],
            ""
        ])
        
        return "\n".join(lines)
    
    def generate_docx(
        self,
        evaluations: List[PlatformAssessment],
        recommendations: List[str],
        title: str = "CMS Evaluation Report"
    ) -> bytes:
        """Generate DOCX report (requires python-docx)."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            
            doc = Document()
            doc.add_heading(title, level=0)
            doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")
            
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph("Evaluation of CMS platforms against requirements.")
            
            doc.add_heading("Platform Assessments", level=1)
            for eval in evaluations:
                doc.add_heading(eval.platform, level=2)
                doc.add_paragraph(f"Overall Fit: {eval.overall_fit_score:.2f}/1.0")
                
                doc.add_paragraph("Strengths:")
                for s in eval.strengths:
                    doc.add_paragraph(s, style="List Bullet")
                
                doc.add_paragraph("Weaknesses:")
                for w in eval.weaknesses:
                    doc.add_paragraph(w, style="List Bullet")
                
                doc.add_paragraph(f"Best for: {eval.best_for_use_case}")
            
            doc.add_heading("Recommendations", level=1)
            for r in recommendations:
                doc.add_paragraph(r, style="List Bullet")
            
            # Save to bytes
            import io
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        
        except ImportError:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")


# ============================================================================
# IMPORTS FOR CONVENIENCE
# ============================================================================

from datetime import datetime

__all__ = [
    "CMSOntology",
    "PlatformAssessment",
    "PlatformEvaluator",
    "CapabilityScorer",
    "ReportGenerator"
]
